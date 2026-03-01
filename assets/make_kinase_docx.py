#!/usr/bin/env python3
"""
生成激酶数据表 Word 文档。需要先安装: pip install python-docx
运行: python3 make_kinase_docx.py
"""
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)

ROWS = [
    ("KYC002", "CDK9", "CDK9", "100", "1000"),
    ("KYC002", "CHEK1", "CHEK1", "81", "1000"),
    ("KYC002", "CLK1", "CLK1", "71", "1000"),
    ("KYC002", "CLK2", "CLK2", "61", "1000"),
    ("KYC002", "CLK4", "CLK4", "55", "1000"),
    ("KYC002", "CSF1R", "CSF1R", "100", "1000"),
    ("KYC002", "CSNK1A1", "CSNK1A1", "85", "1000"),
    ("KYC002", "DAPK1", "DAPK1", "74", "1000"),
    ("KYC002", "EGFR", "EGFR", "100", "1000"),
    ("KYC002", "EGFR(L747-E749del, A750P)", "EGFR", "84", "1000"),
    ("KYC002", "EGFR(L858R)", "EGFR", "93", "1000"),
    ("KYC002", "EPHA1", "EPHA1", "100", "1000"),
    ("KYC002", "ERBB2", "ERBB2", "74", "1000"),
    ("KYC002", "ERK1", "MAPK3", "94", "1000"),
    ("KYC002", "ERK2", "MAPK1", "100", "1000"),
    ("KYC002", "ERK8", "MAPK15", "79", "1000"),
    ("KYC002", "FAK", "PTK2", "100", "1000"),
    ("KYC002", "FGFR2", "FGFR2", "99", "1000"),
    ("KYC002", "FLT1", "FLT1", "93", "1000"),
    ("KYC002", "FLT3", "FLT3", "100", "1000"),
    ("KYC002", "FLT3(ITD)", "FLT3", "62", "1000"),
    ("KYC002", "FLT4", "FLT4", "100", "1000"),
    ("KYC002", "FYN", "FYN", "100", "1000"),
    ("KYC002", "GSK3A", "GSK3A", "13", "1000"),
    ("KYC002", "GSK3B", "GSK3B", "41", "1000"),
    ("KYC002", "HCK", "HCK", "99", "1000"),
    ("KYC002", "IGF1R", "IGF1R", "96", "1000"),
    ("KYC002", "IKK-beta", "IKBKB", "60", "1000"),
    ("KYC002", "INSR", "INSR", "74", "1000"),
    ("KYC002", "ITK", "ITK", "97", "1000"),
    ("KYC002", "JAK2(JH1domain-catalytic)", "JAK2", "63", "1000"),
    ("KYC002", "JNK1", "MAPK8", "70", "1000"),
    ("KYC002", "KIT", "KIT", "100", "1000"),
    ("KYC002", "LCK", "LCK", "93", "1000"),
    ("KYC002", "LIMK1", "LIMK1", "97", "1000"),
    ("KYC002", "LIMK2", "LIMK2", "100", "1000"),
    ("KYC002", "LTK", "LTK", "100", "1000"),
    ("KYC002", "LYN", "LYN", "100", "1000"),
    ("KYC002", "LZK", "MAP3K13", "100", "1000"),
    ("KYC002", "MAP3K1", "MAP3K1", "78", "1000"),
    ("KYC002", "MEK1", "MAP2K1", "71", "1000"),
    ("KYC002", "MET", "MET", "93", "1000"),
    ("KYC002", "MST1R", "MST1R", "100", "1000"),
    ("KYC002", "MTOR", "MTOR", "100", "1000"),
    ("KYC002", "MYLK", "MYLK", "53", "1000"),
    ("KYC002", "NLK", "NLK", "100", "1000"),
    ("KYC002", "p38-alpha", "MAPK14", "100", "1000"),
    ("KYC002", "PAK1", "PAK1", "97", "1000"),
    ("KYC002", "PDGFRA", "PDGFRA", "84", "1000"),
    ("G5", "KIT", "KIT", "82", "1000"),
    ("G5", "LCK", "LCK", "93", "1000"),
    ("G5", "LIMK1", "LIMK1", "98", "1000"),
    ("G5", "LIMK2", "LIMK2", "94", "1000"),
    ("G5", "LTK", "LTK", "92", "1000"),
    ("G5", "LYN", "LYN", "98", "1000"),
    ("G5", "LZK", "MAP3K13", "100", "1000"),
    ("G5", "MAP3K1", "MAP3K1", "69", "1000"),
    ("G5", "MEK1", "MAP2K1", "80", "1000"),
    ("G5", "MET", "MET", "83", "1000"),
    ("G5", "MST1R", "MST1R", "100", "1000"),
    ("G5", "MTOR", "MTOR", "100", "1000"),
    ("G5", "MYLK", "MYLK", "46", "1000"),
    ("G5", "NLK", "NLK", "89", "1000"),
    ("G5", "p38-alpha", "MAPK14", "99", "1000"),
    ("G5", "PAK1", "PAK1", "100", "1000"),
    ("G5", "PDGFRA", "PDGFRA", "86", "1000"),
    ("G5", "PIK3CA", "PIK3CA", "92", "1000"),
    ("G5", "PIK3CB", "PIK3CB", "65", "1000"),
    ("G5", "PIK3CD", "PIK3CD", "64", "1000"),
    ("G5", "PIM1", "PIM1", "1.2", "1000"),
    ("G5", "PIM2", "PIM2", "5.2", "1000"),
    ("G5", "PIM3", "PIM3", "1.8", "1000"),
    ("G5", "PKAC-alpha", "PRKACA", "100", "1000"),
    ("G5", "PLK1", "PLK1", "77", "1000"),
    ("G5", "PRKCD", "PRKCD", "55", "1000"),
    ("G5", "PRKD3", "PRKD3", "19", "1000"),
    ("G5", "RET", "RET", "99", "1000"),
    ("G5", "ROS1", "ROS1", "68", "1000"),
    ("G5", "RSK1(Kin.Dom. 1-N-terminal)", "RPS6KA1", "100", "1000"),
    ("G5", "S6K1", "RPS6KB1", "68", "1000"),
    ("G5", "SRC", "SRC", "94", "1000"),
    ("G5", "SYK", "SYK", "49", "1000"),
    ("G5", "TIE1", "TIE1", "99", "1000"),
    ("G5", "TRKA", "NTRK1", "68", "1000"),
    ("KYC002", "ABL1-phosphorylated", "ABL1", "60", "1000"),
    ("KYC002", "AKT1", "AKT1", "100", "1000"),
    ("KYC002", "ALK", "ALK", "94", "1000"),
    ("KYC002", "AURKA", "AURKA", "90", "1000"),
    ("KYC002", "AURKB", "AURKB", "71", "1000"),
    ("KYC002", "AURKC", "AURKC", "87", "1000"),
    ("KYC002", "AXL", "AXL", "98", "1000"),
    ("KYC002", "BLK", "BLK", "91", "1000"),
    ("KYC002", "BRAF", "BRAF", "74", "1000"),
    ("KYC002", "BRAF(V600E)", "BRAF", "13", "1000"),
    ("KYC002", "BTK", "BTK", "100", "1000"),
    ("KYC002", "CAMK1", "CAMK1", "99", "1000"),
    ("KYC002", "CDK2", "CDK2", "100", "1000"),
    ("KYC002", "CDK4-cyclinD1", "CDK4", "95", "1000"),
    ("KYC002", "PIK3CA", "PIK3CA", "86", "1000"),
    ("KYC002", "PIK3CB", "PIK3CB", "85", "1000"),
    ("KYC002", "PIK3CD", "PIK3CD", "98", "1000"),
    ("KYC002", "PIM1", "PIM1", "1.2", "1000"),
    ("KYC002", "PIM2", "PIM2", "4.1", "1000"),
    ("KYC002", "PIM3", "PIM3", "1.5", "1000"),
    ("KYC002", "PKAC-alpha", "PRKACA", "92", "1000"),
    ("KYC002", "PLK1", "PLK1", "68", "1000"),
    ("KYC002", "PRKCD", "PRKCD", "63", "1000"),
    ("KYC002", "PRKD3", "PRKD3", "57", "1000"),
    ("KYC002", "RET", "RET", "100", "1000"),
    ("KYC002", "ROS1", "ROS1", "98", "1000"),
    ("KYC002", "RSK1(Kin.Dom.1-N-terminal)", "RPS6KA1", "100", "1000"),
    ("KYC002", "S6K1", "RPS6KB1", "76", "1000"),
    ("KYC002", "SRC", "SRC", "93", "1000"),
    ("KYC002", "SYK", "SYK", "81", "1000"),
    ("KYC002", "TIE1", "TIE1", "98", "1000"),
    ("KYC002", "TRKA", "NTRK1", "76", "1000"),
]

def main():
    doc = Document()
    doc.add_heading("激酶数据表", 0)
    table = doc.add_table(rows=1 + len(ROWS), cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["ID Prefix", "Gene/Protein 1", "Gene/Protein 2", "Value 1", "Value 2"]):
        hdr[i].text = text
    for row_idx, row_data in enumerate(ROWS):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            row.cells[col_idx].text = cell_text
    out_path = "kinase_table.docx"
    doc.save(out_path)
    print(f"已生成: {out_path}")

if __name__ == "__main__":
    main()
